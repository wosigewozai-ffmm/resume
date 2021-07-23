import pdfplumber
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox
from pdfminer.pdfpage import PDFTextExtractionNotAllowed

# for num in range(1, 7):
#     print(num)
#     path = str(num)+".pdf"
#     try:
#
#
#         # 用文件对象来创建一个pdf文档分析器
#         parser = PDFParser(open(path, 'rb'))
#         # 创建一个PDF文档
#         doc = PDFDocument(parser)
#         # 连接分析器 与文档对象
#         parser.set_document(doc)
#
#         # 提供初始化密码
#
#
#         # 检测文档是否提供txt转换，不提供就忽略
#         if not doc.is_extractable:
#             raise PDFTextExtractionNotAllowed
#         else:
#             # 创建PDf 资源管理器 来管理共享资源
#             rsrcmgr = PDFResourceManager()
#             # 创建一个PDF设备对象
#             laparams = LAParams()
#             device = PDFPageAggregator(rsrcmgr, laparams=laparams)
#             # 创建一个PDF解释器对象
#             interpreter = PDFPageInterpreter(rsrcmgr, device)
#
#             # 循环遍历列表，每次处理一个page的内容
#             for page in PDFPage.create_pages(doc):
#                 interpreter.process_page(page)
#                 # 接受该页面的LTPage对象
#                 layout = device.get_result()
#                 # 这里layout是一个LTPage对象，里面存放着这个 page 解析出的各种对象
#                 # 包括 LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等
#                 for x in layout:
#                     if isinstance(x, LTTextBox):
#                         print(x.get_text().strip())
#
#     except KeyError:
#         with pdfplumber.open(path) as pdf:
#             page_count = len(pdf.pages)
#             print(page_count)  # 得到页数
#             for page in pdf.pages:
#                 print('---------- 第[%d]页 ----------' % page.page_number)
#                 # 获取当前页面的全部文本信息，包括表格中的文字
#                 print(page.extract_text())

from docx import Document
import re
import jieba
import jieba.posseg as psg
import pandas as pd

allAcademy = []
jieba.load_userdict("schooldict.txt")
jieba.load_userdict("academydict.txt")
jieba.load_userdict("Englishdict.txt")
for line in open("schooldict.txt", "r", encoding='utf-8', errors='ignore'):  # 设置文件对象并读取每一行文件
    strline = str(line)
    allAcademy.append(strline[0:strline.find(" ")])

for line in open("academydict.txt", "r", encoding='utf-8', errors='ignore'):  # 设置文件对象并读取每一行文件
    strline = str(line)
    allAcademy.append(strline[0:strline.find(" ")])

doc = Document('5.docx')

# 学校列表
academyList = []
academyPos = []
wordPos = 0
educateFlag = []
educateFlagPos = []
educateNum = 0
academyDistance = 10

# 时间列表
dataList = []
gapList = []
dateList = []
datePoint = 0
datePos = []
gapPos = []
dataPos = []
dateDistance = 10

#  默认简历中出现的第一个姓名为简历者的姓名
findName = False

resultInfo = []  # 提取的结果集

# 每一段的内容
for para in doc.paragraphs:
    str_ = para.text
    str_ = " " + str_ + " "

    #  正则匹配提取时间段
    patter_timeGap = re.compile('\d{4}[\.|\-|/|年]\d{2}[\.|\-|/|月\d{2}]*?[日]*?[\s]*?[-|~][\s]*?\d{4}[\.|\-|/|年]\d{2}[\.|\-|/|月\d{2}]*?[日]*?')
    Gaps = patter_timeGap.findall(str_)
    if len(Gaps) > 0:
        for gap in Gaps:
            # print(gap)
            gapList.append(gap)
            dateList.append(gap)
    patter_timeGaps = re.compile('\d{4}[\.|\-|/|年]\d{2}[\.|-|/|月\d{2}]*?[日]*?[\s]*?[-|~][\s]*?至今')
    gaps = patter_timeGaps.findall(str_)
    Gaps.append(gaps)
    if len(gaps) > 0:
        for gap in gaps:
            # print(gap)
            gapList.append(gap)
            dateList.append(gap)

    #  正则匹配提取所有日期
    patter_date = re.compile(
        '(\d{4}\-\d{2}[\-\d{2}]*?)|(\d{4}年\d{2}月[\d{2}日]*?)|(\d{4}\.\d{2}[\.\d{2}]*?)|(\d{4}/\d{2}[/\d{2}]*?)')
    Dates = patter_date.findall(str_)
    if len(Dates) > 0:
        for Date in Dates:
            date = Date[0] + Date[1] + Date[2] + Date[3]
            dateFlag = True
            for gap in Gaps:
                if str(gap).find(str(date)) != -1:
                    dateFlag = False
            if dateFlag:
                # print(date)
                dataList.append(date)
                dateList.append(date)

    # 正则匹配姓名信息
    pattern_name = re.compile('姓[\s]*名[^\u4E00-\u9FFF]*([\u4E00-\u9FFF]*)[^\u4E00-\u9FFF]')
    person_name = pattern_name.findall(str_)
    if len(person_name) > 0:
        print("姓名：", person_name[0])
        temp = ["姓名", person_name[0]]
        resultInfo.append(temp)
        findName = True

    # 正则匹配身份证号
    pattern_idNumber = re.compile('\D[0-9]{17}[0-9|x|X]\D')
    id_Number = pattern_idNumber.findall(str_)
    if len(id_Number) > 0:
        id_Number[0] = id_Number[0][1:len(id_Number[0]) - 1]
        print("身份证号：", id_Number[0])
        temp = ["身份证号", id_Number[0]]
        resultInfo.append(temp)

    # 正则匹配手机号
    pattern_phoneNumber = re.compile('\D([0-9]{3})[-]*?([0-9]{4})[-]*?([0-9]{4})\D')
    phone_Number = pattern_phoneNumber.findall(str_)
    if len(phone_Number) > 0:
        print("手机号：", phone_Number[0][0] + phone_Number[0][1] + phone_Number[0][2])
        temp = ["手机号", phone_Number[0][0] + phone_Number[0][1] + phone_Number[0][2]]
        resultInfo.append(temp)

    # 正则匹配邮箱信息
    pattern_mailAddress = re.compile('[0-9a-zA-Z.]+@[0-9a-zA-Z.]+?com')
    mail_Address = pattern_mailAddress.findall(str_)
    if len(mail_Address) > 0:
        print("邮箱：", mail_Address[0])
        temp = ["邮箱", mail_Address[0]]
        resultInfo.append(temp)

    # 性别信息
    pattern_sex = re.compile('[^\u4E00-\u9FFF][男|女][^\u4E00-\u9FFF]')
    person_sex = pattern_sex.findall(str_)
    if len(person_sex) > 0:
        person_sex[0] = person_sex[0][1:len(person_sex[0]) - 1]
        print("性别：", person_sex[0])
        temp = ["性别", person_sex[0]]
        resultInfo.append(temp)

    # 正则匹配年龄信息
    pattern_age = re.compile('\D([0-9]{1,3})岁')
    person_age = pattern_age.findall(str_)
    if len(person_age) > 0:
        print("年龄：", person_age[0])
        temp = ["年龄", person_age[0]]
        resultInfo.append(temp)

    # 正则匹配籍贯信息
    pattern_nativePlace = re.compile('籍[\s]*贯[^\u4E00-\u9FFF]*([\u4E00-\u9FFF]*)[^\u4E00-\u9FFF]')
    native_Place = pattern_nativePlace.findall(str_)
    if len(native_Place) > 0:
        print("籍贯：", native_Place[0])
        temp = ["籍贯", native_Place[0]]
        resultInfo.append(temp)

    # 正则匹配民族信息
    pattern_Folk = re.compile('民[\s]*族[^\u4E00-\u9FFF]*([\u4E00-\u9FFF]*)[^\u4E00-\u9FFF]')
    person_Folk = pattern_Folk.findall(str_)
    if len(person_Folk) > 0:
        print("民族：", person_Folk[0])
        temp = ["民族", person_Folk[0]]
        resultInfo.append(temp)
    #  将段落分词
    temp = str_.split()
    str_ = "".join(temp)
    seg_list = jieba.lcut(str_, cut_all=False)
    for seg in seg_list:
        type_seg = psg.cut(seg).__next__().flag
        if type_seg == "nr":
            #  如果还没有找到姓名属性
            if not findName:
                print("姓名：", seg)
                temp = ["姓名", seg]
                resultInfo.append(temp)
                findName = True
        if type_seg == "es":
            print("英语技能:", seg)
            pattern_englishSkill = re.compile(seg + '[\D]*([\d]{1,3})[\D]')
            english_Skill = pattern_englishSkill.findall(str_)
            if len(english_Skill) > 0:
                temp_skill = "{" + seg + "," + english_Skill[0] + "}"
            else:
                temp_skill = "{" + seg + "," + "None" + "}"
            temp = ["英语技能", temp_skill]
            resultInfo.append(temp)
        wordPos = wordPos + 1

        # 查找日期地址
        if len(dateList) > datePoint:
            if str(dateList[datePoint]).find(seg) == 0:
                datePos.append(wordPos)
                if (len(dateList[datePoint]) > 10) | (str(dateList[datePoint]).find("至今") != -1):
                    gapPos.append(wordPos)
                else:
                    dataPos.append(wordPos)
                # print(wordPos, " ", dateList[datePoint])
                datePoint = datePoint + 1

        # 是否存在XX大学/学院
        if seg in allAcademy:
            academyList.append(seg)
            academyPos.append(wordPos)
        if (seg == "专科") | (seg == "本科") | (seg == "硕士") | (seg == "博士"):
            if not (seg in educateFlag):
                educateNum = educateNum + 1
            educateFlag.append(seg)
            educateFlagPos.append(wordPos)

# 不妨假设第一个单独出现的日期是出生日期
if len(dataList) > 0:
    print("出生日期:", dataList[0])
    resultInfo.append(["出生日期", dataList[0]])

# 匹配学历信息
for i in range(0, len(academyList) - educateNum + 1):
    for j in range(0, len(educateFlag) - educateNum + 1):
        flag = True
        duplicate = []
        for k in range(0, educateNum):
            if educateFlag[j + k] in duplicate:
                flag = False
            duplicate.append(educateFlag[j + k])
            if abs(academyPos[i + k] - educateFlagPos[j + k]) > academyDistance:
                flag = False
        if flag:
            dateStart = 0
            for t in range(0, len(gapList) - educateNum + 1):
                dateFlag = True
                for s in range(0, educateNum):
                    if abs(academyPos[i + s] - gapPos[t + s]) > dateDistance:
                        dateFlag = False
                if dateFlag:
                    dateStart = t
            for k in range(0, educateNum):
                print(educateFlag[j + k], ":", academyList[i + k], gapList[dateStart + k])
                temp = [educateFlag[j + k], academyList[i + k], gapList[dateStart + k]]
                resultInfo.append(temp)

    #  输出到csv表
    result = pd.DataFrame(data=resultInfo)
    result.to_csv('result.csv', encoding="utf_8_sig")
